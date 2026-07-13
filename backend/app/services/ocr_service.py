"""Document text extraction: PyMuPDF for digital PDFs, PaddleOCR for scans/images,
python-docx and openpyxl for Office formats. pdfplumber for table extraction."""
import logging

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

_paddle = None


def _get_paddle():
    global _paddle
    if _paddle is None:
        from paddleocr import PaddleOCR
        _paddle = PaddleOCR(use_angle_cls=True, lang="en", show_log=False)
    return _paddle


def is_scanned_pdf(path: str) -> bool:
    """Heuristic: if the first 3 pages yield <50 chars of embedded text, treat as scanned."""
    doc = fitz.open(path)
    text = "".join(doc[i].get_text() for i in range(min(3, len(doc))))
    doc.close()
    return len(text.strip()) < 50


def extract_pdf_digital(path: str) -> list[dict]:
    """Returns [{page, text}] preserving page numbers."""
    doc = fitz.open(path)
    pages = [{"page": i + 1, "text": doc[i].get_text()} for i in range(len(doc))]
    doc.close()
    return pages


def extract_pdf_scanned(path: str) -> list[dict]:
    """Rasterize each page and run PaddleOCR."""
    import numpy as np
    ocr = _get_paddle()
    doc = fitz.open(path)
    pages = []
    for i in range(len(doc)):
        pix = doc[i].get_pixmap(dpi=200)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
        if pix.n == 4:
            img = img[:, :, :3]
        result = ocr.ocr(img, cls=True)
        lines = []
        for block in result or []:
            for line in block or []:
                lines.append(line[1][0])
        pages.append({"page": i + 1, "text": "\n".join(lines)})
    doc.close()
    return pages


def extract_image(path: str) -> list[dict]:
    ocr = _get_paddle()
    result = ocr.ocr(path, cls=True)
    lines = []
    for block in result or []:
        for line in block or []:
            lines.append(line[1][0])
    return [{"page": 1, "text": "\n".join(lines)}]


def extract_docx(path: str) -> list[dict]:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return [{"page": 1, "text": "\n".join(parts)}]


def extract_xlsx(path: str) -> list[dict]:
    """Each sheet becomes a markdown table + natural-language sentences per row."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    pages = []
    for si, sheet in enumerate(wb.worksheets):
        rows = [[("" if c is None else str(c)) for c in row] for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if any(x.strip() for x in r)]
        if not rows:
            continue
        header, body = rows[0], rows[1:]
        md = "| " + " | ".join(header) + " |\n"
        sentences = []
        for r in body:
            md += "| " + " | ".join(r) + " |\n"
            pairs = [f"{h}: {v}" for h, v in zip(header, r) if v.strip()]
            if pairs:
                sentences.append(f"Record — {'; '.join(pairs)}.")
        pages.append({"page": si + 1, "text": f"Sheet: {sheet.title}\n{md}\n" + "\n".join(sentences)})
    return pages


def extract_pdf_tables(path: str) -> list[dict]:
    """pdfplumber table extraction; returns [{page, table(list of rows)}]."""
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            for table in page.extract_tables() or []:
                out.append({"page": i + 1, "table": table})
    return out


def table_to_sentences(table: list[list]) -> str:
    """Convert an extracted table (first row = header) to natural language for embedding."""
    if not table or len(table) < 2:
        return ""
    header = [(h or "").strip() for h in table[0]]
    sentences = []
    for row in table[1:]:
        pairs = [f"{h} is {v}" for h, v in zip(header, row) if v and str(v).strip()]
        if pairs:
            sentences.append(", ".join(pairs) + ".")
    return " ".join(sentences)


def extract_any(path: str, original_name: str) -> dict:
    """Dispatch by extension. Returns {pages: [{page,text}], tables: [...], method: str}."""
    name = original_name.lower()
    if name.endswith(".pdf"):
        if is_scanned_pdf(path):
            return {"pages": extract_pdf_scanned(path), "tables": [], "method": "paddleocr"}
        return {"pages": extract_pdf_digital(path), "tables": extract_pdf_tables(path), "method": "pymupdf"}
    if name.endswith(".docx"):
        return {"pages": extract_docx(path), "tables": [], "method": "python-docx"}
    if name.endswith((".xlsx", ".xls")):
        return {"pages": extract_xlsx(path), "tables": [], "method": "openpyxl"}
    if name.endswith((".jpg", ".jpeg", ".png")):
        return {"pages": extract_image(path), "tables": [], "method": "paddleocr"}
    if name.endswith((".txt", ".md")):
        with open(path, encoding="utf-8", errors="replace") as f:
            return {"pages": [{"page": 1, "text": f.read()}], "tables": [], "method": "plaintext"}
    raise ValueError(f"Unsupported file type: {original_name}")
